"""
Транскрибер — Аудио → speech2text.ru → LLM → Word (DOCX)
"""

import os
import sys
import io
import json
import time
import re
import hashlib
import logging
import subprocess
import threading
import queue
import requests
from pathlib import Path
from datetime import datetime

# Принудительно UTF-8 для вывода в консоль Windows
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")


# ═══ Логирование ════════════════════════════════════════════════

def setup_logging(log_dir: Path):
    log_path = log_dir / "transcriber.log"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s  %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
        handlers=[
            logging.FileHandler(log_path, encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )

def log(msg: str):
    logging.info(msg)
    # звуковой сигнал при завершении
    if "[ГОТОВО]" in msg:
        print("\a", end="", flush=True)

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ═══ Настройки ══════════════════════════════════════════════════

COLOR_DOCTOR = RGBColor(0x1F, 0x4E, 0x79)
COLOR_PATIENT = RGBColor(0x7B, 0x2D, 0x26)
COLOR_NOTE = RGBColor(0x2E, 0x5E, 0x1F)
COLOR_TIMESTAMP = RGBColor(0x88, 0x88, 0x88)
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)
COLOR_HEADER = RGBColor(0x1F, 0x29, 0x37)
COLOR_BLUE_H = RGBColor(0x1F, 0x4E, 0x79)


# ═══ Конфиг ═════════════════════════════════════════════════════

def load_config() -> dict:
    config_path = Path(__file__).parent / "config.json"
    if not config_path.exists():
        print(f"[ОШИБКА] config.json не найден: {config_path}")
        sys.exit(1)
    with open(config_path, "r", encoding="utf-8") as f:
        return json.load(f)


# ═══ Трекер ═════════════════════════════════════════════════════

# ─────────────────────────────────────────────────────────────────
# Очередь-ожидание при нехватке минут speech2text.
# Если минуты STT кончились — НЕ роняем файл, а ставим обработку на паузу,
# дёшево проверяем баланс и ждём пополнения; при пополнении продолжаем.
# ─────────────────────────────────────────────────────────────────
class TokensExhausted(Exception):
    """Минуты/токены STT закончились — файл НЕ ошибка, ждём пополнения."""
    pass


_WAIT = {"paused": False, "paused_at_min": None}
WAIT_RESUME_MIN = 5         # минимум доступных минут, чтобы возобновить обработку
WAIT_RESUME_MARGIN = 5      # и при этом баланс должен вырасти на столько выше уровня паузы
WAIT_POLL_SECONDS = 120     # как часто проверять пополнение (дёшево, без запуска STT)
WAIT_NOTIFY_COOLDOWN = 1800 # антиспам: один и тот же статус в TG не чаще раза в 30 мин
_LAST_NOTIFY: dict = {}     # kind -> time.monotonic() последней отправки статуса


def stt_minutes_available(config: dict):
    """Сколько минут доступно на speech2text (без запуска распознавания). None при ошибке."""
    try:
        key = config.get("stt_api_key", "")
        if not key:
            return None
        r = requests.get("https://speech2text.ru/api/user/amounts",
                         params={"api-key": key}, timeout=20)
        data = r.json()
        return float(data.get("minutes", {}).get("available", 0))
    except Exception:
        return None


def _tg_notify(config: dict, text: str):
    """Статус оператору в Telegram (через прокси). Best-effort: молча пропускаем без настроек."""
    try:
        token = config.get("telegram_bot_token", "")
        chat = config.get("operator_chat_id", "")
        if not token or not chat:
            return
        proxy = config.get("telegram_proxy", "") or None
        proxies = {"http": proxy, "https": proxy} if proxy else None
        requests.post(f"https://api.telegram.org/bot{token}/sendMessage",
                      data={"chat_id": chat, "text": text},
                      proxies=proxies, timeout=30)
    except Exception:
        pass


def _notify_status(config: dict, kind: str, text: str):
    """TG-уведомление о статусе паузы с антиспам-кулдауном (на каждый kind отдельно).
    Защищает от спама, даже если цикл по какой-то причине снова зафлапает."""
    now = time.monotonic()
    if now - _LAST_NOTIFY.get(kind, -1e9) < WAIT_NOTIFY_COOLDOWN:
        return
    _LAST_NOTIFY[kind] = now
    _tg_notify(config, text)


def _enter_wait_mode(config: dict, queue_size: int):
    if _WAIT["paused"]:
        return
    _WAIT["paused"] = True
    # Запоминаем уровень минут на момент паузы: возобновим только когда баланс
    # реально вырастет выше него (а не на застрявшем остатке вроде ~4 мин).
    _WAIT["paused_at_min"] = stt_minutes_available(config)
    log(f"  [ОЖИДАНИЕ ТОКЕНОВ] Минуты speech2text закончились — {queue_size} в очереди, жду пополнения.")
    _notify_status(config, "exhausted",
                   f"⏳ Закончились минуты speech2text. Записей в очереди: {queue_size}. "
                   f"Они НЕ потеряны — пополните минуты на speech2text.ru, и обработка продолжится автоматически.")


def _exit_wait_mode(config: dict, avail: float):
    if not _WAIT["paused"]:
        return
    _WAIT["paused"] = False
    _WAIT["paused_at_min"] = None
    log(f"  [ВОЗОБНОВЛЕНИЕ] Минуты пополнены (~{avail:.0f}) — продолжаю обработку очереди.")
    _notify_status(config, "resumed",
                   f"✅ Минуты speech2text пополнены (~{avail:.0f}). Продолжаю обработку очереди.")


def load_processed(path: Path) -> set:
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            return set(line.strip() for line in f if line.strip())
    return set()


def save_processed(path: Path, processed: set):
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(sorted(processed)))


# ═══ gen-api.ru STT (ElevenLabs / Silero) ═══════════════════════

STT_ENDPOINTS = {
    "elevenlabs": "https://api.gen-api.ru/api/v1/networks/speech-to-text",
    "silero": "https://api.gen-api.ru/api/v1/networks/silero-stt",
    "nemotron": "https://api.gen-api.ru/api/v1/networks/nemotron-asr",
}


def send_to_stt_sync(audio_path: str, api_key: str, lang: str = "ru", stt_model: str = "silero") -> str:
    log(f"  [STT] Отправляю: {Path(audio_path).name} (модель: {stt_model})")

    size_mb = os.path.getsize(audio_path) / (1024 * 1024)
    log(f"  [STT] Размер: {size_mb:.1f} МБ")

    url = STT_ENDPOINTS.get(stt_model, STT_ENDPOINTS["silero"])
    headers = {
        "Accept": "application/json",
        "Authorization": f"Bearer {api_key}",
    }

    import mimetypes
    mime = mimetypes.guess_type(audio_path)[0] or "audio/mpeg"

    with open(audio_path, "rb") as audio_file:
        if stt_model == "elevenlabs":
            data = {
                "is_sync": "true",
                "diarize": "true",
                "tag_audio_events": "true",
                "model": "v2",
            }
            if lang:
                data["language_code"] = lang
            files = {"audio_url": (Path(audio_path).name, audio_file, mime)}
        elif stt_model == "nemotron":
            data = {"is_sync": "true"}
            files = {"audio_url": (Path(audio_path).name, audio_file, mime)}
        else:
            data = {"is_sync": "true"}
            files = {"payload": (Path(audio_path).name, audio_file, mime)}

        log(f"  [STT] Отправляю файл ({size_mb:.1f} МБ), синхронный режим...")
        resp = requests.post(url, headers=headers, files=files, data=data, timeout=600)

    if resp.status_code not in (200, 201):
        raise Exception(f"STT ошибка {resp.status_code}: {resp.text}")

    result = resp.json()
    log(f"  [STT] Ключи ответа: {list(result.keys())}")

    output = result.get("response") or result.get("output")
    if not output:
        raise Exception(f"Пустой ответ STT: {json.dumps(result, ensure_ascii=False)[:500]}")

    text = _extract_text_from_output(output)
    log(f"  [STT] Получено: {len(text)} символов")
    return text


def _format_time(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    if h > 0:
        return f"{h}:{m:02d}:{s:02d}"
    return f"{m}:{s:02d}"


def _extract_text_from_output(output) -> str:
    """Извлекает текст из output gen-api.ru (ElevenLabs формат)."""
    if isinstance(output, str):
        return output.strip()

    if isinstance(output, dict):
        words = output.get("words")
        if words and isinstance(words, list) and len(words) > 0 and isinstance(words[0], dict) and "speaker_id" in words[0]:
            return _words_to_diarized_text(words)

        utterances = output.get("utterances") or output.get("segments")
        if utterances:
            lines = []
            for u in utterances:
                speaker = u.get("speaker_id") or u.get("speaker", "")
                text = u.get("text", "").strip()
                if speaker:
                    lines.append(f"{speaker}: {text}")
                else:
                    lines.append(text)
            return "\n".join(lines)

        text = output.get("text") or output.get("transcript", "")
        return str(text).strip()

    if isinstance(output, list):
        parts = []
        for item in output:
            if isinstance(item, dict):
                parts.append(item.get("text", ""))
            else:
                parts.append(str(item))
        return "\n".join(p for p in parts if p).strip()

    return str(output).strip()


def _words_to_diarized_text(words: list) -> str:
    """Группирует words по спикерам и формирует текст с таймкодами."""
    segments = []
    current_speaker = None
    current_text = []
    current_start = 0.0

    for w in words:
        if w.get("type") != "word":
            continue
        speaker = w.get("speaker_id", "")
        text = w.get("text", "")

        if speaker != current_speaker:
            if current_speaker is not None and current_text:
                segments.append((current_speaker, current_start, " ".join(current_text)))
            current_speaker = speaker
            current_text = [text]
            current_start = w.get("start", 0.0)
        else:
            current_text.append(text)

    if current_speaker is not None and current_text:
        segments.append((current_speaker, current_start, " ".join(current_text)))

    speaker_map = {}
    counter = 1
    lines = []
    for speaker_id, start, text in segments:
        if speaker_id not in speaker_map:
            speaker_map[speaker_id] = f"Спикер {counter}"
            counter += 1
        label = speaker_map[speaker_id]
        ts = _format_time(start)
        lines.append(f"{label} [{ts}]: {text}")

    return "\n".join(lines)


# ═══ speech2text.ru STT (async polling) ═════════════════════════

S2T_BASE = "https://speech2text.ru/api"


def send_to_speech2text(audio_path: str, api_key: str, lang: str, speakers,
                         poll_interval: int, max_wait: int) -> str:
    log(f"  [STT] Отправляю: {Path(audio_path).name} (speech2text.ru)")
    size_mb = os.path.getsize(audio_path) / (1024 * 1024)
    log(f"  [STT] Размер: {size_mb:.1f} МБ")

    # 1. Создаём задачу
    submit_url = f"{S2T_BASE}/recognitions/task/file"
    params = {"api-key": api_key}
    data = {}
    if lang:
        data["lang"] = lang
    if speakers:
        data["speakers"] = speakers

    with open(audio_path, "rb") as audio_file:
        files = {"file": (Path(audio_path).name, audio_file)}
        log(f"  [STT] Создаю задачу...")
        resp = requests.post(submit_url, params=params, files=files, data=data,
                             headers={"Accept": "application/json"}, timeout=600)

    if resp.status_code not in (200, 201):
        raise Exception(f"STT submit ошибка {resp.status_code}: {resp.text[:500]}")

    task = resp.json()
    task_id = task.get("id") or task.get("task_id")
    if not task_id:
        raise Exception(f"Нет task_id в ответе: {json.dumps(task, ensure_ascii=False)[:500]}")
    log(f"  [STT] Задача: {task_id}")

    # 2. Polling
    status_url = f"{S2T_BASE}/recognitions/{task_id}"
    start = time.time()
    last_code = None
    while True:
        elapsed = time.time() - start
        if elapsed > max_wait:
            raise Exception(f"STT превысил max_wait={max_wait}с (последний код: {last_code})")
        time.sleep(poll_interval)
        try:
            r = requests.get(status_url, params=params, timeout=30)
        except requests.RequestException as e:
            log(f"  [STT] Ошибка опроса: {e}, повтор...")
            continue
        if r.status_code != 200:
            log(f"  [STT] Опрос вернул {r.status_code}, повтор...")
            continue
        info = r.json()
        status = info.get("status", {}) or {}
        code = status.get("code")
        desc = status.get("description", "")
        last_code = code
        log(f"  [STT] Статус: {code} ({desc}, {int(elapsed)}с)")
        if code == 200:
            break
        if code == 501:
            raise Exception(f"STT ошибка распознавания: {json.dumps(info, ensure_ascii=False)[:500]}")
        if code == 102:
            # Лимит минут speech2text исчерпан — это НЕ ошибка файла. Специальное
            # исключение: главный цикл поставит очередь на паузу и дождётся
            # пополнения, ничего не потеряв (см. TokensExhausted / wait-режим).
            raise TokensExhausted(desc or "лимит минут исчерпан")

    # 3. Результат (txt — простой текст)
    result_url = f"{S2T_BASE}/recognitions/{task_id}/result/txt"
    r = requests.get(result_url, params=params, timeout=120)
    if r.status_code != 200:
        raise Exception(f"Не удалось получить результат: {r.status_code} {r.text[:500]}")

    text = r.text.strip()
    if not text:
        raise Exception("Пустой результат от speech2text.ru")

    log(f"  [STT] Получено: {len(text)} символов")
    return text


def transcribe_audio(audio_path: str, config: dict) -> str:
    api_key = config["stt_api_key"]
    stt_model = config.get("stt_model", "silero")
    lang = config.get("stt_lang", "ru")

    if stt_model == "speech2text":
        return send_to_speech2text(
            audio_path, api_key, lang,
            config.get("stt_speakers"),
            config.get("stt_poll_interval", 10),
            config.get("stt_max_wait", 1800),
        )
    return send_to_stt_sync(audio_path, api_key, lang, stt_model)


# ═══ LLM ═════════════════════════════════════════════════════

def build_system_prompt(prompt_template: str, config: dict) -> str:
    variants = config.get("doctor_names_stt_variants", [])
    variants_str = ", ".join(variants) if variants else "нет"
    replacements = {
        "{clinic_name}": config.get("clinic_name", "Клиника"),
        "{doctor_name}": config.get("doctor_name", "Врач"),
        "{doctor_full}": config.get("doctor_full", config.get("doctor_name", "Врач")),
        "{doctor_specialty}": config.get("doctor_specialty", ""),
        "{doctor_names_stt_variants}": variants_str,
    }
    result = prompt_template
    for key, value in replacements.items():
        result = result.replace(key, value)
    return result


def _extract_json_block(text: str) -> str | None:
    """Извлекает первый сбалансированный {...} блок, учитывая строки и экранирование."""
    start = text.find('{')
    if start < 0:
        return None
    depth = 0
    in_str = False
    escape = False
    for i in range(start, len(text)):
        c = text[i]
        if escape:
            escape = False
            continue
        if c == '\\':
            escape = True
            continue
        if c == '"':
            in_str = not in_str
            continue
        if in_str:
            continue
        if c == '{':
            depth += 1
        elif c == '}':
            depth -= 1
            if depth == 0:
                return text[start:i + 1]
    return None


# ═══ Нормализатор вложенного Calgary-Cambridge JSON в плоский ════

def flatten_summary(summary: dict) -> dict:
    """Преобразует вложенный Calgary-Cambridge summary (полный промт) в плоский (haiku промт).
    Если уже плоский — возвращает как есть."""
    if not isinstance(summary, dict):
        return summary or {}
    # already flat? — heuristic: top-level field that's only in flat schema
    if "presenting_complaint" in summary or "anamnesis" in summary:
        return summary

    out: dict = {
        "patient": summary.get("patient"),
        "doctor": summary.get("doctor"),
        "doctor_specialty": summary.get("doctor_specialty"),
        "referring_doctor": summary.get("referring_doctor"),
    }

    init = summary.get("initiating_the_session") or {}
    out["presenting_complaint"] = init.get("presenting_complaint")
    out["agenda"] = init.get("agenda")

    gathering = summary.get("gathering_information") or {}
    disease = gathering.get("disease_framework") or {}
    anamnesis_parts = []
    for k in ("history_of_present_illness", "obstetric_gynecologic_history", "menstrual_cycle",
              "reproductive_plans", "past_medical_history", "family_history", "previous_investigations"):
        v = disease.get(k)
        if v:
            anamnesis_parts.append(str(v))
    out["anamnesis"] = "\n\n".join(anamnesis_parts) if anamnesis_parts else None

    illness = gathering.get("illness_framework") or {}
    out["patient_ideas"] = illness.get("ideas")
    out["patient_concerns"] = illness.get("concerns")
    out["patient_expectations"] = illness.get("expectations")
    out["effect_on_life"] = illness.get("effect_on_life")
    feelings = illness.get("feelings")
    if feelings:
        cur = out.get("patient_concerns")
        out["patient_concerns"] = (cur + f"\nЭмоции: {feelings}") if cur else f"Эмоции: {feelings}"

    exam = summary.get("examination_and_investigations") or {}
    findings_parts = []
    if exam.get("physical_examination"):
        findings_parts.append(f"Физикальный осмотр: {exam['physical_examination']}")
    us = exam.get("ultrasound_findings") or {}
    if us:
        us_lines = []
        for k, label in [
            ("uterus", "Матка"), ("endometrium", "Эндометрий"), ("cervix", "Шейка матки"),
            ("right_ovary", "Правый яичник"), ("left_ovary", "Левый яичник"),
            ("pouch_of_douglas", "Позадиматочное пространство"), ("additional", "Дополнительно"),
        ]:
            v = us.get(k)
            if v:
                us_lines.append(f"{label}: {v}")
        if us_lines:
            findings_parts.append("УЗИ ОМТ:\n" + "\n".join(us_lines))
    if exam.get("other_imaging"):
        findings_parts.append(f"Другие исследования: {exam['other_imaging']}")
    out["findings"] = "\n\n".join(findings_parts) if findings_parts else None

    out["labs_and_markers"] = exam.get("labs_and_markers")
    out["histology"] = exam.get("histology")

    plan = summary.get("explanation_and_planning") or {}
    out["diagnosis_explained"] = plan.get("diagnosis_explained")
    out["treatment_options"] = plan.get("treatment_options")
    surgery = plan.get("surgery_logistics")
    if isinstance(surgery, dict):
        parts = []
        for k, label in [
            ("approach", "Доступ"), ("duration", "Длительность"),
            ("anesthesia", "Наркоз"), ("hospitalization", "Госпитализация"),
            ("recovery", "Восстановление"),
        ]:
            v = surgery.get(k)
            if v:
                parts.append(f"{label}: {v}")
        out["surgery_logistics"] = "\n".join(parts) if parts else None
    else:
        out["surgery_logistics"] = surgery
    out["cost"] = plan.get("cost")

    closing = summary.get("closing_the_session") or {}
    out["decision"] = closing.get("decision")
    out["next_steps"] = closing.get("next_steps")
    out["safety_netting"] = closing.get("safety_netting")
    out["follow_up"] = closing.get("follow_up")

    # Новые поля (медицинская запись поверх Calgary-Cambridge):
    # demographics, icd10_codes, red_flags, missing_or_to_clarify.
    # В full-промте они на том же top-level summary, что и Calgary-Cambridge секции.
    out["demographics"] = summary.get("demographics")
    out["icd10_codes"] = summary.get("icd10_codes") or []
    out["red_flags"] = summary.get("red_flags") or []
    out["missing_or_to_clarify"] = summary.get("missing_or_to_clarify") or []

    return out


def normalize_response(data: dict) -> dict:
    """Применяет flatten_summary к каждой консультации в data. Idempotent."""
    if not isinstance(data, dict):
        return data
    consultations = data.get("consultations")
    if isinstance(consultations, list):
        for c in consultations:
            if isinstance(c, dict) and "summary" in c:
                c["summary"] = flatten_summary(c.get("summary") or {})
    return data


def _extract_balanced_from(text: str, start: int) -> str | None:
    """Balanced {...} match starting at given index (text[start] must be '{')."""
    if start < 0 or start >= len(text) or text[start] != '{':
        return None
    depth = 0
    in_str = False
    escape = False
    for i in range(start, len(text)):
        c = text[i]
        if escape:
            escape = False; continue
        if c == '\\':
            escape = True; continue
        if c == '"':
            in_str = not in_str; continue
        if in_str:
            continue
        if c == '{':
            depth += 1
        elif c == '}':
            depth -= 1
            if depth == 0:
                return text[start:i + 1]
    return None


def parse_json_output(output: str) -> dict | None:
    """Робастный парсер.
    Стратегия:
    1) direct json.loads
    2) ищет сбалансированный {...} блок, СОДЕРЖАЩИЙ ключ "consultations"
       (выбирает САМЫЙ ВНЕШНИЙ такой блок — устойчиво к prose-before-JSON,
       и не путает с вложенным фрагментом dialog при усечённом начале)
    3) fallback на первый {...} блок
    Всегда сохраняет raw output в _llm_debug.txt для последующей диагностики."""
    text = output.strip()
    text = re.sub(r'^```(?:json)?\s*', '', text)
    text = re.sub(r'\s*```$', '', text)
    text = text.strip()

    # Always save raw for diagnostics
    debug_path = Path(__file__).parent / "temp" / "_llm_debug.txt"
    debug_path.parent.mkdir(exist_ok=True)
    try:
        with open(debug_path, "w", encoding="utf-8") as f:
            f.write(output)
    except OSError:
        pass

    # 1. Direct
    try:
        data = json.loads(text)
        log("  [LLM] JSON распарсен")
        return data
    except json.JSONDecodeError:
        pass

    # 2. Find OUTERMOST block that contains "consultations" key
    anchor_pos = text.find('"consultations"')
    if anchor_pos >= 0:
        # walk backwards from anchor looking for the OUTERMOST { that wraps it
        # (the earliest { whose balanced match includes the anchor)
        best_block = None
        for start in range(anchor_pos - 1, -1, -1):
            if text[start] == '{':
                block = _extract_balanced_from(text, start)
                if block and '"consultations"' in block:
                    # remember this candidate; keep going to find earlier (outer) wrapper
                    best_block = block
        if best_block:
            try:
                data = json.loads(best_block)
                log(f"  [LLM] JSON распарсен (внешний блок с consultations, {len(best_block)} символов)")
                return data
            except json.JSONDecodeError as e:
                log(f"  [LLM] Блок с consultations не парсится: {e}")

    # 3. Fallback: first balanced block
    extracted = _extract_json_block(text)
    if extracted:
        try:
            data = json.loads(extracted)
            log(f"  [LLM] JSON распарсен (fallback, первый блок {len(extracted)} символов)")
            return data
        except json.JSONDecodeError as e:
            log(f"  [LLM] Fallback-блок не парсится: {e}")
    else:
        log("  [LLM] Сбалансированный JSON-блок не найден в ответе")

    log(f"  [LLM] Сырой ответ доступен: temp/{debug_path.name}")
    return None


def process_with_llm(raw_text: str, system_prompt: str, config: dict,
                         model_override: str = None, file_label: str = "",
                         effort: str = "high", max_tokens: int = 16000):
    """Вызывает языковую модель напрямую по HTTP (Messages-style API:
    POST {base_url}/v1/messages). Без CLI-обёрток — только requests.

    base_url / api_key берутся из config (llm_base_url / llm_api_key),
    с fallback на env LLM_BASE_URL / LLM_API_KEY. Подходит любой LLM-провайдер
    или шлюз с совместимым Messages-style HTTP API.

    model_override позволяет роутеру выбрать модель не из config.
    file_label — имя обрабатываемого файла, попадает в [LLM]-логи (нужно для бот-парсера
    в параллельном режиме, чтобы различать какой файл сейчас в обработке).

    Возвращает dict с разобранным JSON (успех), либо сырую строку-ответ
    (если JSON не распарсился — _llm_only сделает robust-parse/fallback),
    либо None (если API так и не ответил).
    """
    model = model_override or config.get("llm_model", "your-llm-model")
    label = f"«{file_label}» " if file_label else ""
    log(f"  [LLM] Обрабатываю файл {label}({len(raw_text)} символов) через {model} (API)...")

    retry_count = config.get("llm_retry_count", 3)
    # Таймаут ОДНОГО запроса. llm_timeout (часы) — это бюджет на весь файл,
    # для отдельного HTTP-запроса берём llm_request_timeout (по умолчанию 900с).
    timeout_sec = config.get("llm_request_timeout", 900)
    out_tokens = config.get("llm_max_tokens", max_tokens)

    base_url = (config.get("llm_base_url")
                or os.environ.get("LLM_BASE_URL")
                or "").rstrip("/")
    api_key = config.get("llm_api_key") or os.environ.get("LLM_API_KEY")
    if not base_url:
        log("  [LLM] НЕТ llm_base_url (ни в config, ни в env) — не могу вызвать API")
        return None
    if not api_key:
        log("  [LLM] НЕТ llm_api_key (ни в config, ни в env) — не могу вызвать API")
        return None
    url = f"{base_url}/v1/messages"

    # Сохраняем системный промт для диагностики (как раньше)
    prompt_tmp = Path(__file__).parent / "temp" / "_system_prompt.txt"
    prompt_tmp.parent.mkdir(exist_ok=True)
    try:
        with open(prompt_tmp, "w", encoding="utf-8") as f:
            f.write(system_prompt)
    except OSError:
        pass

    # Sandwich the raw text between instruction reminders.
    # At very long inputs, models lose top-of-prompt instructions; reinforce at the end.
    user_message = (
        "Сырая расшифровка ниже. Верни СТРОГО валидный JSON по заданной схеме — "
        "никакого текста, цитат или объяснений до или после. Без markdown-обёрток "
        "(никаких ```json``` тегов).\n\n"
        "ОБЯЗАТЕЛЬНО: твой ответ ДОЛЖЕН начинаться буквально с символов:\n"
        '`{"consultations": [`\n'
        "и заканчиваться `]}`. Не начинай с dialog-entries или summary напрямую — "
        "сначала ВНЕШНЯЯ обёртка с consultations.\n\n"
        "=== НАЧАЛО РАСШИФРОВКИ ===\n\n"
        f"{raw_text}\n\n"
        "=== КОНЕЦ РАСШИФРОВКИ ===\n\n"
        'Напоминание ещё раз: ответ начинается с `{"consultations": [` и далее по схеме.'
    )

    headers = {
        "x-api-key": api_key,
        "content-type": "application/json",
    }
    body = {
        "model": model,
        "max_tokens": out_tokens,
        "system": system_prompt,
        "messages": [{"role": "user", "content": user_message}],
    }

    # Статусы, которые имеет смысл повторить
    RETRYABLE = {408, 409, 425, 429, 500, 502, 503, 504, 529}
    last_output = None

    for attempt in range(1, retry_count + 1):
        if attempt > 1:
            log(f"  [LLM] Попытка {attempt}/{retry_count}...")
        try:
            resp = requests.post(url, headers=headers, json=body, timeout=timeout_sec)

            if resp.status_code != 200:
                snippet = (resp.text or "")[:500]
                log(f"  [LLM] HTTP {resp.status_code}: {snippet}")
                if resp.status_code in RETRYABLE and attempt < retry_count:
                    time.sleep(min(60, 5 * attempt))
                    continue
                if attempt == retry_count:
                    return last_output
                time.sleep(5)
                continue

            j = resp.json()
            output = "".join(
                b.get("text", "") for b in (j.get("content") or [])
                if isinstance(b, dict) and b.get("type") == "text"
            ).strip()
            stop_reason = j.get("stop_reason")
            usage = j.get("usage") or {}
            log(f"  [LLM] Ответ: {len(output)} символов "
                f"(stop={stop_reason}, in={usage.get('input_tokens')}, out={usage.get('output_tokens')})")

            if not output:
                raise Exception("API вернул пустой content")
            last_output = output

            if stop_reason == "max_tokens":
                log(f"  [LLM] ⚠ ответ обрезан по max_tokens={out_tokens} — JSON может быть неполным")

            data = parse_json_output(output)
            if data is not None:
                return data

            log("  [LLM] JSON невалидный, повторяю...")

        except requests.Timeout:
            log(f"  [LLM] Таймаут запроса (попытка {attempt}, {timeout_sec}s)")
        except Exception as e:
            log(f"  [LLM] Ошибка (попытка {attempt}): {type(e).__name__}: {e}")

        if attempt == retry_count:
            # Вернём сырой текст последнего ответа (если был) — _llm_only
            # сделает robust-parse и при неудаче — fallback-docx с расшифровкой.
            return last_output
        time.sleep(5)

    log("  [LLM] Все попытки исчерпаны")
    return last_output


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_HEADER
        p.space_before = Pt(24)
        p.space_after = Pt(12)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_BLUE_H
        p.space_before = Pt(20)
        p.space_after = Pt(8)
    return p


def add_summary_field(doc, label: str, value):
    if value is None:
        return
    if isinstance(value, list):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}:")
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_HEADER
        for item in value:
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"  • {item}")
            run2.font.name = "Arial"
            run2.font.size = Pt(11)
            p2.paragraph_format.left_indent = Cm(1)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_HEADER
        run2 = p.add_run(str(value))
        run2.font.name = "Arial"
        run2.font.size = Pt(11)


def add_dialog_line(doc, entry: dict):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    speaker = entry.get("speaker", "")
    time_str = entry.get("time", "")
    text = entry.get("text", "")

    if time_str:
        run_time = p.add_run(f"[{time_str}]  ")
        run_time.font.name = "Consolas"
        run_time.font.size = Pt(9)
        run_time.font.color.rgb = COLOR_TIMESTAMP

    if speaker == "doctor":
        run_sp = p.add_run("Врач: ")
        run_sp.font.color.rgb = COLOR_DOCTOR
        run_sp.font.bold = True
        run_sp.font.name = "Arial"
        run_sp.font.size = Pt(11)
        run_t = p.add_run(text)
        run_t.font.name = "Arial"
        run_t.font.size = Pt(11)

    elif speaker == "patient":
        run_sp = p.add_run("Пациентка: ")
        run_sp.font.color.rgb = COLOR_PATIENT
        run_sp.font.bold = True
        run_sp.font.name = "Arial"
        run_sp.font.size = Pt(11)
        run_t = p.add_run(text)
        run_t.font.name = "Arial"
        run_t.font.size = Pt(11)

    elif speaker == "note":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_n = p.add_run(f"[ {text} ]")
        run_n.font.name = "Arial"
        run_n.font.size = Pt(10)
        run_n.font.color.rgb = COLOR_NOTE
        run_n.font.italic = True

    else:
        run_t = p.add_run(f"{speaker}: {text}")
        run_t.font.name = "Arial"
        run_t.font.size = Pt(11)


def generate_docx(consultation: dict, output_path: Path, record_name: str = "", config: dict = None):
    if config is None:
        config = {}
    doc = Document()

    section = doc.sections[0]
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    summary = consultation.get("summary", {})
    patient_name = consultation.get("patient_name", "Пациентка")
    duration = consultation.get("duration", "")

    clinic = config.get("clinic_name", "Клиника")
    # Специальность — из определённой моделью по содержанию приёма (врач может быть
    # любой: репродуктолог, хирург и т.д.), а НЕ из дефолта config.
    specialty = (summary.get("doctor_specialty") or "").strip()
    title = f"{clinic} — {specialty}" if specialty else clinic

    # Заголовок
    add_heading_styled(doc, title, level=1)

    subtitle = "Расшифровка аудиозаписи приёма"
    if duration:
        subtitle += f" ({duration})"
    if record_name:
        subtitle += f"  •  {record_name}"
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(subtitle)
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = COLOR_GRAY
    run_sub.font.italic = True

    # Разделитель
    p_sep = doc.add_paragraph()
    run_sep = p_sep.add_run("─" * 70)
    run_sep.font.color.rgb = COLOR_GRAY
    run_sep.font.size = Pt(8)

    # Шапка: участники
    fields_head = [
        ("Пациентка", summary.get("patient")),
        ("Врач", summary.get("doctor")),
        ("Направивший врач", summary.get("referring_doctor")),
    ]
    for label, value in fields_head:
        add_summary_field(doc, label, value)

    # Демография (если что-то заполнено)
    demo = summary.get("demographics") or {}
    if isinstance(demo, dict) and any(demo.get(k) for k in
                                       ("age", "bmi", "pregnancy", "allergies", "current_medications")):
        add_heading_styled(doc, "Демография и текущий статус", level=2)
        add_summary_field(doc, "Возраст", demo.get("age"))
        add_summary_field(doc, "ИМТ", demo.get("bmi"))
        add_summary_field(doc, "Беременность", demo.get("pregnancy"))
        add_summary_field(doc, "Аллергии", demo.get("allergies"))
        add_summary_field(doc, "Текущие препараты", demo.get("current_medications"))

    # 🚩 КРАСНЫЕ ФЛАГИ — поднимаем НАВЕРХ docx, перед стадиями Calgary-Cambridge.
    # Если зафиксированы — выделяем красным, это критично для медицинской безопасности.
    red_flags = summary.get("red_flags") or []
    # Фильтруем placeholder-инструкции которые модель могла оставить из шаблона
    red_flags = [
        rf for rf in red_flags
        if isinstance(rf, str) and rf.strip()
        and not rf.lower().startswith(("список", "беременность —", "гинекология —"))
    ]
    if red_flags:
        add_heading_styled(doc, "🚩 КРАСНЫЕ ФЛАГИ", level=2)
        for flag in red_flags:
            p_rf = doc.add_paragraph()
            p_rf.paragraph_format.left_indent = Cm(0.5)
            rl = p_rf.add_run("🚩 "); rl.font.bold = True; rl.font.color.rgb = COLOR_PATIENT
            rl.font.name = "Arial"; rl.font.size = Pt(11)
            rv = p_rf.add_run(str(flag))
            rv.font.name = "Arial"; rv.font.size = Pt(11)
            rv.font.color.rgb = COLOR_PATIENT

    # === 1. Начало консультации ===
    if summary.get("presenting_complaint") or summary.get("agenda"):
        add_heading_styled(doc, "1. Начало консультации", level=2)
        add_summary_field(doc, "Главная жалоба", summary.get("presenting_complaint"))
        add_summary_field(doc, "Повестка приёма", summary.get("agenda"))

    # === 2. Сбор анамнеза ===
    if any(summary.get(k) for k in ("anamnesis", "patient_ideas", "patient_concerns",
                                     "patient_expectations", "effect_on_life")):
        add_heading_styled(doc, "2. Сбор анамнеза", level=2)
        add_summary_field(doc, "Анамнез", summary.get("anamnesis"))
        if any(summary.get(k) for k in ("patient_ideas", "patient_concerns",
                                         "patient_expectations", "effect_on_life")):
            # ICE-блок
            p_ice = doc.add_paragraph()
            run_ice = p_ice.add_run("Перспектива пациентки (ICE):")
            run_ice.font.name = "Arial"
            run_ice.font.size = Pt(11)
            run_ice.font.bold = True
            run_ice.font.color.rgb = COLOR_PATIENT
            add_summary_field(doc, "  Представления (Ideas)", summary.get("patient_ideas"))
            add_summary_field(doc, "  Опасения (Concerns)", summary.get("patient_concerns"))
            add_summary_field(doc, "  Ожидания (Expectations)", summary.get("patient_expectations"))
            add_summary_field(doc, "  Влияние на жизнь", summary.get("effect_on_life"))

    # === 3. Осмотр и исследования ===
    if any(summary.get(k) for k in ("findings", "labs_and_markers", "histology")):
        add_heading_styled(doc, "3. Осмотр и разбор исследований", level=2)
        add_summary_field(doc, "Находки УЗИ/МРТ", summary.get("findings"))
        add_summary_field(doc, "Лаб. показатели и онкомаркеры", summary.get("labs_and_markers"))
        add_summary_field(doc, "Гистология / цитология", summary.get("histology"))

    # === 4. Объяснение и планирование ===
    if any(summary.get(k) for k in ("diagnosis_explained", "treatment_options",
                                     "surgery_logistics", "cost", "icd10_codes")):
        add_heading_styled(doc, "4. Объяснение и планирование", level=2)
        add_summary_field(doc, "Объяснение диагноза", summary.get("diagnosis_explained"))
        # МКБ-10 коды если врач их озвучил
        icd_codes = summary.get("icd10_codes") or []
        icd_codes = [c for c in icd_codes if isinstance(c, str) and c.strip()
                     and not c.lower().startswith("мкб-10 код")]  # фильтр placeholder'а
        if icd_codes:
            add_summary_field(doc, "МКБ-10", icd_codes)
        opts = summary.get("treatment_options")
        if isinstance(opts, list) and opts:
            p_t = doc.add_paragraph()
            run_t = p_t.add_run("Варианты лечения:")
            run_t.font.name = "Arial"
            run_t.font.size = Pt(11)
            run_t.font.bold = True
            run_t.font.color.rgb = COLOR_HEADER
            for opt in opts:
                if isinstance(opt, dict):
                    p_opt = doc.add_paragraph()
                    p_opt.paragraph_format.left_indent = Cm(0.5)
                    run_n = p_opt.add_run(f"• {opt.get('option', '')}")
                    run_n.font.name = "Arial"
                    run_n.font.size = Pt(11)
                    run_n.font.bold = True
                    if opt.get("description"):
                        p_d = doc.add_paragraph()
                        p_d.paragraph_format.left_indent = Cm(1)
                        run_d = p_d.add_run(opt["description"])
                        run_d.font.name = "Arial"
                        run_d.font.size = Pt(11)
                    if opt.get("pros"):
                        p_p = doc.add_paragraph()
                        p_p.paragraph_format.left_indent = Cm(1)
                        rl = p_p.add_run("+ "); rl.font.bold = True; rl.font.color.rgb = COLOR_NOTE
                        rl.font.name = "Arial"; rl.font.size = Pt(11)
                        rv = p_p.add_run(str(opt["pros"]))
                        rv.font.name = "Arial"; rv.font.size = Pt(11)
                    if opt.get("cons"):
                        p_c = doc.add_paragraph()
                        p_c.paragraph_format.left_indent = Cm(1)
                        rl = p_c.add_run("− "); rl.font.bold = True; rl.font.color.rgb = COLOR_PATIENT
                        rl.font.name = "Arial"; rl.font.size = Pt(11)
                        rv = p_c.add_run(str(opt["cons"]))
                        rv.font.name = "Arial"; rv.font.size = Pt(11)
                    if opt.get("indications_for_patient"):
                        p_i = doc.add_paragraph()
                        p_i.paragraph_format.left_indent = Cm(1)
                        rl = p_i.add_run("Для пациентки: "); rl.font.bold = True
                        rl.font.name = "Arial"; rl.font.size = Pt(11)
                        rv = p_i.add_run(str(opt["indications_for_patient"]))
                        rv.font.name = "Arial"; rv.font.size = Pt(11)
                else:
                    add_summary_field(doc, "  •", str(opt))
        elif opts:
            add_summary_field(doc, "Варианты лечения", str(opts))
        add_summary_field(doc, "Логистика операции", summary.get("surgery_logistics"))
        add_summary_field(doc, "Стоимость", summary.get("cost"))

    # === 5. Завершение приёма ===
    if any(summary.get(k) for k in ("decision", "next_steps", "safety_netting", "follow_up")):
        add_heading_styled(doc, "5. Завершение приёма", level=2)
        add_summary_field(doc, "Итоговое решение / тактика", summary.get("decision"))
        add_summary_field(doc, "Дальнейшие шаги", summary.get("next_steps"))
        # Safety-netting — выделим, это важно
        sn = summary.get("safety_netting")
        if sn:
            p_sn = doc.add_paragraph()
            rl = p_sn.add_run("⚠ Что должно насторожить: ")
            rl.font.bold = True; rl.font.color.rgb = COLOR_PATIENT
            rl.font.name = "Arial"; rl.font.size = Pt(11)
            rv = p_sn.add_run(str(sn))
            rv.font.name = "Arial"; rv.font.size = Pt(11)
        add_summary_field(doc, "Следующий контакт", summary.get("follow_up"))

    # === 📋 Что уточнить на следующем приёме ===
    missing = summary.get("missing_or_to_clarify") or []
    missing = [
        m for m in missing
        if isinstance(m, str) and m.strip()
        and not m.lower().startswith(("список", "что не хватает"))  # фильтр placeholder'а
    ]
    if missing:
        add_heading_styled(doc, "📋 Что уточнить на следующем приёме", level=2)
        for item in missing:
            p_m = doc.add_paragraph()
            p_m.paragraph_format.left_indent = Cm(0.5)
            run_m = p_m.add_run(f"☐  {item}")
            run_m.font.name = "Arial"; run_m.font.size = Pt(11)

    # Диалог рендерим только если есть содержимое
    # (в новой схеме LLM не выводит dialog — он идёт отдельным файлом-расшифровкой)
    dialog_entries = consultation.get("dialog") or []
    if dialog_entries:
        # Разделитель
        p_sep2 = doc.add_paragraph()
        run_sep2 = p_sep2.add_run("─" * 70)
        run_sep2.font.color.rgb = COLOR_GRAY
        run_sep2.font.size = Pt(8)

        add_heading_styled(doc, "Диалог", level=2)
        p_legend = doc.add_paragraph()
        r1 = p_legend.add_run("■ Врач")
        r1.font.color.rgb = COLOR_DOCTOR; r1.font.name = "Arial"
        r1.font.size = Pt(10); r1.font.bold = True
        p_legend.add_run("    ")
        r2 = p_legend.add_run("■ Пациентка")
        r2.font.color.rgb = COLOR_PATIENT; r2.font.name = "Arial"
        r2.font.size = Pt(10); r2.font.bold = True
        p_legend.add_run("    ")
        r3 = p_legend.add_run("■ Служебное")
        r3.font.color.rgb = COLOR_NOTE; r3.font.name = "Arial"
        r3.font.size = Pt(10); r3.font.bold = True
        doc.add_paragraph()
        for entry in dialog_entries:
            add_dialog_line(doc, entry)

    doc.save(str(output_path))
    log(f"  [DOCX] Сохранён: {output_path.name}")


def generate_transcript_docx(raw_text: str, output_path: Path, record_name: str = "", config: dict = None):
    """Сохраняет сырую STT-расшифровку как отдельный docx-файл (без обработки LLM).
    Отдаётся врачу как «компаньон» к сводке для self-verification."""
    if config is None:
        config = {}
    doc = Document()
    section = doc.sections[0]
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    style = doc.styles['Normal']
    style.font.name = 'Arial'; style.font.size = Pt(11)

    clinic = config.get("clinic_name", "Клиника")
    add_heading_styled(doc, f"{clinic} — Сырая расшифровка", level=1)
    subtitle = "Автоматическая расшифровка аудио, без обработки LLM"
    if record_name:
        subtitle += f"  •  {record_name}"
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(subtitle)
    run_sub.font.name = "Arial"; run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = COLOR_GRAY; run_sub.font.italic = True

    p_sep = doc.add_paragraph()
    run_sep = p_sep.add_run("─" * 70)
    run_sep.font.color.rgb = COLOR_GRAY; run_sep.font.size = Pt(8)

    # Каждый блок «Спикер N:\n0:HH:MM:SS - текст» → парный абзац
    for block in raw_text.strip().split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph()
        run = p.add_run(block)
        run.font.name = "Arial"; run.font.size = Pt(11)

    doc.save(str(output_path))
    log(f"  [TRANSCRIPT] Сохранён: {output_path.name}")


# ═══ Fallback при пустом ответе LLM ═════════════════════════════

def _parse_raw_dialog(raw_text: str) -> list:
    """Парсит сырую расшифровку speech2text.ru в список dialog-записей.
    Формат: блоки 'Спикер N:\\n0:HH:MM:SS - текст'."""
    entries = []
    pattern = re.compile(
        r'Спикер\s+(\S+?):\s*\n\s*([0-9:]+)\s*[-—]\s*(.+?)(?=\n\s*Спикер|\Z)',
        re.DOTALL,
    )
    for m in pattern.finditer(raw_text):
        sid = m.group(1).strip()
        time_str = m.group(2).strip()
        text = m.group(3).strip()
        # speaker роль не определена → ренденрим как note чтобы было видно «Спикер 1:» / «Спикер 2:»
        entries.append({
            "time": time_str,
            "speaker": "note",
            "text": f"Спикер {sid}: {text}",
        })
    if not entries:
        # Формат не распознан — одним блоком
        entries = [{"time": "", "speaker": "note", "text": raw_text}]
    return entries


def _build_fallback_consultation(raw_text: str, err_msg: str, config: dict) -> dict:
    """Минимальный консультация-объект с сырой расшифровкой когда LLM отказался
    или вернул пустой результат. Врач получит docx с raw-текстом для ручной обработки."""
    note = err_msg or "LLM не выделил структуру консультации"
    return {
        "patient_name": "Не определено",
        "duration": "",
        "summary": {
            "patient": None,
            "doctor": None,  # врач берётся из расшифровки; по умолчанию не подставляем
            "decision": f"⚠ {note}. Ниже — сырая расшифровка для ручной проверки.",
            "next_steps": "Проверить расшифровку и оформить вручную",
        },
        "dialog": _parse_raw_dialog(raw_text),
    }


# ═══ Пайплайн ══════════════════════════════════════════════════

def file_hash(path: Path) -> str:
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def load_hashes(path: Path) -> set:
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            return set(line.strip() for line in f if line.strip())
    return set()


def save_hashes(path: Path, hashes: set):
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(sorted(hashes)))


def _stt_only(audio_path: Path, config: dict) -> str:
    """STT с кешем. Возвращает raw_text. Бросает Exception при ошибке."""
    filename = audio_path.stem
    temp_folder = Path(__file__).parent / "temp"
    temp_folder.mkdir(exist_ok=True)
    raw_path = temp_folder / f"{filename}_raw.txt"
    if raw_path.exists() and raw_path.stat().st_size > 0:
        with open(raw_path, "r", encoding="utf-8") as f:
            raw_text = f.read()
        log(f"  [STT] Используем кеш: {raw_path.name} ({len(raw_text)} символов)")
        return raw_text
    raw_text = transcribe_audio(str(audio_path), config)
    with open(raw_path, "w", encoding="utf-8") as f:
        f.write(raw_text)
    log(f"  [СОХРАНЕНО] Сырая: {raw_path.name}")
    return raw_text


def _llm_only(audio_path: Path, raw_text: str, output_folder: Path, config: dict) -> bool:
    """LLM + DOCX + cleanup. Раздельно от STT для параллельной обработки.
    Возвращает True/False (False — LLM совсем не вернул данные)."""
    filename = audio_path.stem
    temp_folder = Path(__file__).parent / "temp"
    temp_folder.mkdir(exist_ok=True)

    # Routing: выбираем промт + модель по длине STT-текста
    model_override = None
    if config.get("routing_enabled"):
        threshold = config.get("routing_threshold_chars", 25000)
        prompt_short = config.get("prompt_file_short")
        prompt_long = config.get("prompt_file_long")
        if prompt_short and prompt_long and Path(prompt_short).exists() and Path(prompt_long).exists():
            if len(raw_text) >= threshold:
                prompt_file = Path(prompt_long)
                model_override = config.get("llm_model_long", "your-llm-model")
                log(f"  [РОУТИНГ] «{audio_path.name}»: {len(raw_text)} символов ≥ {threshold} → {model_override} + полный промт")
            else:
                prompt_file = Path(prompt_short)
                model_override = config.get("llm_model_short", "your-llm-model")
                log(f"  [РОУТИНГ] «{audio_path.name}»: {len(raw_text)} символов < {threshold} → {model_override} + короткий промт")
        else:
            prompt_file = Path(config["prompt_file"])
            log(f"  [РОУТИНГ] Файлы промтов не настроены, использую legacy: {prompt_file}")
    else:
        prompt_file = Path(config["prompt_file"])

    if not prompt_file.exists():
        log(f"  [ОШИБКА] Промт не найден: {prompt_file}")
        return True

    with open(prompt_file, "r", encoding="utf-8") as f:
        prompt_template = f.read().strip()
    system_prompt = build_system_prompt(prompt_template, config)

    data = process_with_llm(raw_text, system_prompt, config,
                                model_override=model_override,
                                file_label=audio_path.name)
    if not data:
        log(f"  [!] LLM не вернул данные для «{audio_path.name}»")
        return False

    # process_with_llm now may return either a parsed dict OR a raw string
    # (when the proxy wraps the model output in markdown fences). Normalise here.
    if isinstance(data, str):
        text = data.strip()
        # Use robust parser that extracts a balanced {...} block even when
        # the response is truncated by the proxy at ~20k chars.
        data = parse_json_output(text) or {
            "consultations": [_build_fallback_consultation(text, "no balanced JSON block", config)]
        }
        if isinstance(data, dict) and data.get("consultations"):
            log(f"  [LLM-API] robust-parse recovered {len(data['consultations'])} consultation(s) from {len(text)} chars")
        else:
            log(f"  [LLM-API] robust-parse failed on {len(text)} chars; writing raw fallback")

    data = normalize_response(data)

    # Strip "dialog" arrays from each consultation — the raw transcript is
    # already saved as a separate DOCX, so duplicating it inside the summary
    # just blows the response past the proxy ~21k char cap.
    if isinstance(data, dict) and isinstance(data.get("consultations"), list):
        for c in data["consultations"]:
            if isinstance(c, dict):
                c.pop("dialog", None)
                if isinstance(c.get("summary"), dict):
                    c["summary"].pop("dialog", None)

    json_path = temp_folder / f"{filename}_data.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    consultations = data.get("consultations") or []
    err = data.get("error")
    if err or not consultations:
        log(f"  [!] «{audio_path.name}»: {err or 'Нет консультаций в ответе'} — создаю fallback docx")
        consultations = [_build_fallback_consultation(raw_text, err, config)]

    clinic = config.get("clinic_name", "Клиника").replace(" ", "_")
    if len(consultations) == 1:
        c = consultations[0]
        patient = (c.get("patient_name") or "").replace(" ", "_")
        name = f"{clinic}_{filename}_{patient}.docx" if patient else f"{clinic}_{filename}.docx"
        generate_docx(c, output_folder / name, filename, config)
    else:
        log(f"  [ИНФО] «{audio_path.name}»: {len(consultations)} консультаций")
        for i, c in enumerate(consultations, 1):
            patient = (c.get("patient_name") or "").replace(" ", "_")
            name = f"{clinic}_{filename}_{i}_{patient}.docx" if patient else f"{clinic}_{filename}_{i}.docx"
            generate_docx(c, output_folder / name, f"{filename} (конс. {i})", config)

    # Параллельно сохраняем сырую расшифровку отдельным файлом-компаньоном.
    # Имя совпадает с базовым именем summary-docx + суффикс «_расшифровка», чтобы
    # бот/клиент мог найти оба по одному префиксу.
    transcript_name = f"{clinic}_{filename}_расшифровка.docx"
    try:
        generate_transcript_docx(raw_text, output_folder / transcript_name, filename, config)
    except Exception as e:
        log(f"  [!] не удалось сохранить расшифровку: {type(e).__name__}: {e}")

    # Очистка промежуточных файлов
    for tmp in temp_folder.glob(f"{filename}*"):
        try:
            tmp.unlink()
        except OSError:
            pass

    log(f"  [ГОТОВО] «{audio_path.name}» обработка завершена")
    return True


def process_file(audio_path: Path, output_folder: Path, config: dict, processed_hashes: set, hashes_path: Path) -> bool:
    """Legacy sequential pipeline: STT → LLM → DOCX в одном потоке."""
    filename = audio_path.stem

    import re as _re; _safe_fn = _re.sub(r"[\*\?\[\]]", "_", filename); existing = list(output_folder.glob(f"*{_safe_fn}*.docx"))
    if existing:
        log(f"  [ПРОПУСК] DOCX уже есть: {existing[0].name}")
        return True

    h = file_hash(audio_path)
    if h in processed_hashes:
        # Дубль по хэшу — но если docx уже сожран cleanup'ом, обрабатываем заново.
        # Иначе пользователь, переотправивший файл, никогда не получит результат.
        if any(output_folder.glob(f"*{filename}*.docx")):
            log(f"  [ДУБЛИКАТ] {audio_path.name} — есть готовый docx")
            return True
        log(f"  [ДУБЛИКАТ-RECOVER] {audio_path.name} — хэш есть, docx нет (удалён cleanup'ом) — обработаю заново")
        processed_hashes.discard(h)
        save_hashes(hashes_path, processed_hashes)

    log(f"\n{'='*60}")
    log(f"  Файл: {audio_path.name}")
    log(f"{'='*60}")

    try:
        size1 = audio_path.stat().st_size
        time.sleep(2)
        size2 = audio_path.stat().st_size
        if size1 != size2 or size1 == 0:
            log(f"  [ОЖИДАНИЕ] Файл копируется...")
            return False
    except OSError as e:
        log(f"  [ОШИБКА] {e}")
        return False

    try:
        raw_text = _stt_only(audio_path, config)
    except TokensExhausted:
        raise   # пробрасываем наверх — главный цикл поставит очередь на паузу
    except Exception as e:
        log(f"  [ОШИБКА STT] {e}")
        return False

    return _llm_only(audio_path, raw_text, output_folder, config)


# ═══ Поллинг ════════════════════════════════════════════════════

def main():
    config = load_config()

    watch_folder = Path(config["watch_folder"])
    output_folder = Path(config["output_folder"])
    scan_interval = config.get("scan_interval", 15)
    extensions = set(config.get("audio_extensions", [".mp3", ".wav", ".m4a"]))

    tracker_path = Path(__file__).parent / "processed.txt"
    processed = load_processed(tracker_path)
    hashes_path = Path(__file__).parent / "processed_hashes.txt"
    processed_hashes = load_hashes(hashes_path)
    max_retries = config.get("max_retries", 3)
    fail_counts = {}

    try:
        output_folder.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        print(f"[ОШИБКА] {e}")
        sys.exit(1)

    setup_logging(Path(__file__).parent)

    clinic = config.get("clinic_name", "Транскрибер")
    doctor = config.get("doctor_name", "")

    log("╔════════════════════════════════════════════════════════╗")
    log(f"║  {clinic} — Транскрибер{'':<31}║")
    stt_model = config.get("stt_model", "silero")
    log(f"║  Аудио → speech2text.ru          → LLM API → Word   ║")
    log("╠════════════════════════════════════════════════════════╣")
    log(f"║  Вход:  {str(watch_folder):<46} ║")
    log(f"║  Выход: {str(output_folder):<46} ║")
    log(f"║  Скан:  каждые {scan_interval} сек{'':<32} ║")
    log("╠════════════════════════════════════════════════════════╣")
    log("║  Ctrl+C чтобы остановить                              ║")
    log("╚════════════════════════════════════════════════════════╝")

    if processed:
        log(f"\n[ИНФО] Ранее обработано: {len(processed)}")

    log(f"\n[МОНИТОРИНГ] Слежу за: {watch_folder}")
    log(f"[МОНИТОРИНГ] Кидай аудио — обработаю автоматически\n")

    parallel = config.get("parallel_enabled", True)
    if parallel:
        log(f"[РЕЖИМ] Параллельный (STT в main-потоке, LLM в фоновом)")
        _run_parallel(
            watch_folder, output_folder, config, extensions,
            processed, processed_hashes, tracker_path, hashes_path,
            fail_counts, max_retries, scan_interval,
        )
    else:
        log(f"[РЕЖИМ] Последовательный (legacy)")
        _run_sequential(
            watch_folder, output_folder, config, extensions,
            processed, processed_hashes, tracker_path, hashes_path,
            fail_counts, max_retries, scan_interval,
        )


def _run_sequential(watch_folder, output_folder, config, extensions,
                     processed, processed_hashes, tracker_path, hashes_path,
                     fail_counts, max_retries, scan_interval):
    try:
        while True:
            # Очередь-ожидание при нехватке минут STT (см. _run_parallel).
            if _WAIT["paused"]:
                avail = stt_minutes_available(config)
                base = _WAIT.get("paused_at_min") or 0
                threshold = max(WAIT_RESUME_MIN, base + WAIT_RESUME_MARGIN)
                if avail is not None and avail >= threshold:
                    _exit_wait_mode(config, avail)
                else:
                    time.sleep(WAIT_POLL_SECONDS)
                    continue
            try:
                audio_files = [
                    f for f in watch_folder.iterdir()
                    if f.is_file() and f.suffix.lower() in extensions
                ]
            except OSError as e:
                log(f"[!] Папка недоступна: {e}")
                time.sleep(scan_interval)
                continue

            new_files = [
                f for f in audio_files
                if f.name not in processed and fail_counts.get(f.name, 0) < max_retries
            ]

            if new_files:
                log(f"\n[НАЙДЕНО] {len(new_files)} новых файлов")
                for af in sorted(new_files):
                    try:
                        success = process_file(af, output_folder, config, processed_hashes, hashes_path)
                    except TokensExhausted:
                        _enter_wait_mode(config, len(new_files))
                        break
                    if success:
                        processed.add(af.name)
                        save_processed(tracker_path, processed)
                        h = file_hash(af)
                        processed_hashes.add(h)
                        save_hashes(hashes_path, processed_hashes)
                    else:
                        fail_counts[af.name] = fail_counts.get(af.name, 0) + 1
                        if fail_counts[af.name] >= max_retries:
                            log(f"  [!] {af.name}: {max_retries} неудачных попыток, пропускаю")
                log(f"\n[МОНИТОРИНГ] Продолжаю...\n")

            time.sleep(scan_interval)

    except KeyboardInterrupt:
        log("\n[СТОП] Завершаю...")
        save_processed(tracker_path, processed)
        log(f"Обработано: {len(processed)}. До встречи!")


def _run_parallel(watch_folder, output_folder, config, extensions,
                   processed, processed_hashes, tracker_path, hashes_path,
                   fail_counts, max_retries, scan_interval):
    """Параллельная обработка: STT в main, LLM в worker-потоке.
    Пока LLM думает над файлом A, STT уже идёт по файлу B."""
    llm_queue: "queue.Queue" = queue.Queue()
    in_pipeline: set = set()  # filenames currently in STT or LLM (под state_lock)
    state_lock = threading.Lock()
    shutdown = threading.Event()

    def llm_worker():
        while not shutdown.is_set():
            try:
                item = llm_queue.get(timeout=2)
            except queue.Empty:
                continue
            if item is None:
                return
            audio_path, raw_text, audio_hash = item
            try:
                success = _llm_only(audio_path, raw_text, output_folder, config)
                with state_lock:
                    if success:
                        processed.add(audio_path.name)
                        save_processed(tracker_path, processed)
                        processed_hashes.add(audio_hash)
                        save_hashes(hashes_path, processed_hashes)
                    else:
                        fail_counts[audio_path.name] = fail_counts.get(audio_path.name, 0) + 1
                        if fail_counts[audio_path.name] >= max_retries:
                            log(f"  [!] {audio_path.name}: {max_retries} неудачных попыток (LLM)")
            except Exception as e:
                log(f"  [ОШИБКА LLM-WORKER] {audio_path.name}: {type(e).__name__}: {e}")
                with state_lock:
                    fail_counts[audio_path.name] = fail_counts.get(audio_path.name, 0) + 1
            finally:
                with state_lock:
                    in_pipeline.discard(audio_path.name)

    worker_t = threading.Thread(target=llm_worker, daemon=True, name="llm-worker")
    worker_t.start()

    try:
        while True:
            # Очередь-ожидание: минуты STT кончились — не запускаем распознавание,
            # дёшево проверяем баланс и ждём пополнения; файлы остаются в input/.
            if _WAIT["paused"]:
                avail = stt_minutes_available(config)
                base = _WAIT.get("paused_at_min") or 0
                threshold = max(WAIT_RESUME_MIN, base + WAIT_RESUME_MARGIN)
                if avail is not None and avail >= threshold:
                    _exit_wait_mode(config, avail)
                else:
                    time.sleep(WAIT_POLL_SECONDS)
                    continue
            try:
                audio_files = [
                    f for f in watch_folder.iterdir()
                    if f.is_file() and f.suffix.lower() in extensions
                ]
            except OSError as e:
                log(f"[!] Папка недоступна: {e}")
                time.sleep(scan_interval)
                continue

            with state_lock:
                new_files = [
                    f for f in audio_files
                    if f.name not in processed
                    and f.name not in in_pipeline
                    and fail_counts.get(f.name, 0) < max_retries
                ]

            if new_files:
                log(f"\n[НАЙДЕНО] {len(new_files)} новых файлов (очередь LLM: {llm_queue.qsize()})")

            for af in sorted(new_files):
                # Pre-checks (cheap, в main потоке)
                filename = af.stem
                import re as _re; _safe_fn = _re.sub(r"[\*\?\[\]]", "_", filename); existing = list(output_folder.glob(f"*{_safe_fn}*.docx"))
                if existing:
                    log(f"  [ПРОПУСК] {af.name}: DOCX уже есть ({existing[0].name})")
                    with state_lock:
                        processed.add(af.name)
                        save_processed(tracker_path, processed)
                    continue

                try:
                    h = file_hash(af)
                except OSError as e:
                    log(f"  [ОШИБКА] {af.name}: {e}")
                    continue

                with state_lock:
                    if h in processed_hashes:
                        # Дубль по хэшу — если docx нет (cleanup съел), снимаем хэш и обрабатываем заново
                        has_docx = any(output_folder.glob(f"*{af.stem}*.docx"))
                        if has_docx:
                            log(f"  [ДУБЛИКАТ] {af.name} — есть готовый docx")
                            processed.add(af.name)
                            save_processed(tracker_path, processed)
                            continue
                        log(f"  [ДУБЛИКАТ-RECOVER] {af.name} — хэш есть, docx нет — обработаю заново")
                        processed_hashes.discard(h)
                        save_hashes(hashes_path, processed_hashes)

                # Резервируем слот в pipeline (под lock-ом)
                with state_lock:
                    in_pipeline.add(af.name)

                log(f"\n{'='*60}")
                log(f"  Файл: {af.name}")
                log(f"{'='*60}")

                # Проверка стабильности файла (не копируется ли)
                try:
                    size1 = af.stat().st_size
                    time.sleep(2)
                    size2 = af.stat().st_size
                    if size1 != size2 or size1 == 0:
                        log(f"  [ОЖИДАНИЕ] {af.name}: файл копируется, попробую позже")
                        with state_lock:
                            in_pipeline.discard(af.name)
                        continue
                except OSError as e:
                    log(f"  [ОШИБКА] {af.name}: {e}")
                    with state_lock:
                        in_pipeline.discard(af.name)
                    continue

                # STT (в main потоке — пока LLM может работать над предыдущим файлом)
                try:
                    raw_text = _stt_only(af, config)
                except TokensExhausted:
                    # Минуты STT кончились — НЕ ошибка файла: возвращаем его в очередь,
                    # ставим обработку на паузу и ждём пополнения (см. начало цикла).
                    with state_lock:
                        in_pipeline.discard(af.name)
                    _enter_wait_mode(config, len(new_files))
                    break
                except Exception as e:
                    log(f"  [ОШИБКА STT] {af.name}: {type(e).__name__}: {e}")
                    with state_lock:
                        fail_counts[af.name] = fail_counts.get(af.name, 0) + 1
                        in_pipeline.discard(af.name)
                    continue

                # В очередь к LLM-worker'у — main продолжает сканить дальше
                llm_queue.put((af, raw_text, h))
                log(f"  [QUEUE] «{af.name}» → LLM (глубина очереди: {llm_queue.qsize()})")

            if new_files:
                log(f"\n[МОНИТОРИНГ] Продолжаю...\n")

            time.sleep(scan_interval)

    except KeyboardInterrupt:
        log("\n[СТОП] Завершаю...")
        shutdown.set()
        llm_queue.put(None)
        worker_t.join(timeout=10)
        save_processed(tracker_path, processed)
        log(f"Обработано: {len(processed)}. До встречи!")


if __name__ == "__main__":
    main()
